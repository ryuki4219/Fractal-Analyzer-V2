import os
import json
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def replace_in_runs(paragraph, mapping):
    for run in paragraph.runs:
        for key, val in mapping.items():
            if key in run.text:
                run.text = run.text.replace(key, val)


def remove_todo_paragraphs(document):
    to_remove = []
    for p in document.paragraphs:
        text = p.text.strip()
        if text.startswith("[TODO]"):
            to_remove.append(p)
    for p in to_remove:
        p._element.getparent().remove(p._element)


def insert_section_text_after_heading(document, heading_text, content):
    # Find heading paragraph by exact text match
    for i, p in enumerate(document.paragraphs):
        if p.text.strip() == heading_text:
            # Insert content after heading
            insert_idx = i + 1
            new_p = document.paragraphs[insert_idx]._element
            # Workaround: add a new paragraph at the end and move it
            tail = document.add_paragraph("")
            tail_run = tail.add_run(content)
            # Move element near the target index by XML operations
            p._element.addnext(tail._element)
            return True
    return False


def fill_template(template_path, config_path, output_path):
    doc = Document(template_path)

    with open(config_path, 'r', encoding='utf-8') as f:
        cfg = json.load(f)

    # Global simple replacements for header fields
    mapping = {
        "[TITLE]": cfg.get("title", ""),
        "[AUTHORS]": cfg.get("authors", ""),
        "[CONTACT]": cfg.get("contact", ""),
    }
    for p in doc.paragraphs:
        replace_in_runs(p, mapping)

    # Keywords line
    keywords = cfg.get("keywords")
    if keywords:
        kw_text = ", ".join(keywords)
        for p in doc.paragraphs:
            if p.text.strip().startswith("Keywords:"):
                # Replace entire paragraph text
                p.clear() if hasattr(p, 'clear') else None
                p.text = f"Keywords: {kw_text}"
                break

    # Abstract content (first paragraph after heading "Abstract")
    abstract = cfg.get("abstract")
    if abstract:
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip() == "Abstract":
                # Next non-empty paragraph replaced with abstract
                j = i + 1
                while j < len(doc.paragraphs):
                    if doc.paragraphs[j].text.strip():
                        doc.paragraphs[j].clear() if hasattr(doc.paragraphs[j], 'clear') else None
                        doc.paragraphs[j].text = abstract
                        break
                    j += 1
                break

    # Sections mapping: insert text after each heading if provided
    sections = cfg.get("sections", {})
    for heading, content in sections.items():
        insert_section_text_after_heading(doc, heading, content)

    # Remove [TODO] paragraphs if requested
    if cfg.get("remove_todo_hints", False):
        remove_todo_paragraphs(doc)

    # Ensure output directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)


if __name__ == "__main__":
    base_dir = os.path.dirname(os.path.dirname(__file__))
    template_path = os.path.join(base_dir, "docs", "templates", "Paper_Template.docx")
    config_path = os.path.join(base_dir, "docs", "templates", "paper_config.json")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    output_dir = os.path.join(base_dir, "docs", "output")
    output_path = os.path.join(output_dir, f"Paper_Filled_{timestamp}.docx")

    fill_template(template_path, config_path, output_path)
    print(f"Saved: {output_path}")
