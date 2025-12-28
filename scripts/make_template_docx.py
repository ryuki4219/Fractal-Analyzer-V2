from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE


def add_heading(document, text, level=1):
    document.add_heading(text, level=level)


def add_paragraph(document, text, bold=False, italic=False, align=None):
    p = document.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if align:
        p.alignment = align
    return p


def add_placeholder_section(document, title, level=1, hint=None):
    add_heading(document, title, level)
    if hint:
        add_paragraph(document, f"[TODO] {hint}")
    document.add_paragraph("\n")


def add_table(document, rows, cols, header=None):
    table = document.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    if header:
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(header):
            hdr_cells[i].text = h
    return table


def build_template(output_path):
    doc = Document()

    # Base styles
    styles = doc.styles
    if "Body Text" not in [s.name for s in styles]:
        body_style = styles.add_style("Body Text", WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.name = "Times New Roman"
        body_style._element.rPr.rFonts.set(qn('w:eastAsia'), "Yu Mincho")
        body_style.font.size = Pt(11)

    # Title Page
    p = add_paragraph(doc, "[TITLE] フラクタルを活用した顔画像解析（論文タイトル案）", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.runs[0].font.size = Pt(18)
    add_paragraph(doc, "[AUTHORS] 著者名（所属）", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "[CONTACT] 対応著者メール / 住所 / ORCID", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "\n")

    add_heading(doc, "Abstract", level=1)
    add_paragraph(doc, "[TODO] 200-300語。背景→目的→方法→結果→結論の順で簡潔に。")
    add_paragraph(doc, "Keywords: [TODO] fractal dimension, skin analysis, texture, dermatology")

    # 1. Introduction
    add_placeholder_section(doc, "1. Introduction", level=1, hint="背景、課題設定、貢献（本研究の新規性・意義）を明確化。")

    # 2. Related Work
    add_placeholder_section(doc, "2. Related Work", level=1, hint="皮膚画像解析（GLCM/LBP/Haralick, CNN転移学習）とフラクタルの医用画像応用を整理。")

    # 3. Methods
    add_heading(doc, "3. Methods", level=1)

    add_placeholder_section(doc, "3.1 Dataset", level=2, hint="データ構成、枚数、撮影条件、ラベル分布、匿名化・同意。")
    add_placeholder_section(doc, "3.2 Preprocessing", level=2, hint="肌領域抽出、照明補正、解像度統一、正規化の詳細。")

    add_heading(doc, "3.3 Fractal Feature Extraction", level=2)
    add_paragraph(doc, "[TODO] 箱数え法（Box-counting）の定義と離散近似の説明：")
    add_paragraph(doc, "D = lim_{ε→0} log N(ε) / log(1/ε) を離散スケールに対して線形回帰で推定。")
    add_paragraph(doc, "[TODO] スケール系列、二値化・形態処理、線形領域選定、適合度(R²)の記述。")

    add_placeholder_section(doc, "3.4 Baselines", level=2, hint="GLCM/LBP/Haralick、簡易CNN（転移学習）など比較対象の設定。")
    add_placeholder_section(doc, "3.5 Modeling", level=2, hint="分類（SVM/RandomForest）/ 回帰（Lasso/Ridge）などの学習器と設定。")

    # 4. Experiments
    add_heading(doc, "4. Experiments", level=1)
    add_paragraph(doc, "[TODO] 交差検証（stratified k-fold）、検証セット分離、ハイパラ選定。")
    add_paragraph(doc, "[TODO] 多重比較補正（Benjamini–Hochberg等）、効果量・信頼区間。")
    add_paragraph(doc, "\n")

    # Dataset stats table placeholder
    add_paragraph(doc, "Table 1: Dataset Summary")
    table = add_table(doc, rows=4, cols=4, header=["Subset", "Images", "Conditions", "Notes"])
    table.rows[1].cells[0].text = "Train"
    table.rows[2].cells[0].text = "Validation"
    table.rows[3].cells[0].text = "Test"
    add_paragraph(doc, "\n")

    # 5. Results
    add_heading(doc, "5. Results", level=1)
    add_paragraph(doc, "[TODO] 指標（分類: F1/ROC-AUC、回帰: MAE/RMSE）を平均±標準偏差／CIで提示。")
    add_paragraph(doc, "[TODO] 図：ROC曲線、残差プロット、線形フィット例、箱ひげ図。")

    # 6. Discussion
    add_placeholder_section(doc, "6. Discussion", level=1, hint="効く条件・効かない条件、画像分解能・照明の影響、臨床解釈、限界。")

    # 7. Ethics
    add_placeholder_section(doc, "7. Ethics", level=1, hint="匿名化、同意取得、再利用範囲、バイアス（年齢・性別・肌色）考察。")

    # 8. Reproducibility
    add_heading(doc, "8. Reproducibility", level=1)
    add_paragraph(doc, "[TODO] コード・設定・乱数種の公開。ファイル紐付け：fd_boxcount.py, fractal_app.py, skin_analysis.py, generate_updated_report.py。")

    # 9. Conclusion
    add_placeholder_section(doc, "9. Conclusion", level=1, hint="本研究の要約、臨床・工学的意義、今後の展望。")

    # Acknowledgments, References, Appendix
    add_placeholder_section(doc, "Acknowledgments", level=1, hint="研究資金や協力者への謝辞。")
    add_heading(doc, "References", level=1)
    add_paragraph(doc, "[TODO] 書式（例：Vancouver/APA）。文献マネージャ連携を推奨。")
    add_placeholder_section(doc, "Appendix", level=1, hint="追加図・ハイパーパラメータ・失敗例・感度分析。")

    # Footer note
    add_paragraph(doc, "\n")
    add_paragraph(doc, "Note: このテンプレートはPDF骨子に基づく論文化用の初期雛形です。章・図表は必要に応じて調整してください。", italic=True)

    doc.save(output_path)


if __name__ == "__main__":
    import os
    out_dir = os.path.join("docs", "templates")
    os.makedirs(out_dir, exist_ok=True)
    output_path = os.path.join(out_dir, "Paper_Template.docx")
    build_template(output_path)
    print(f"Saved: {output_path}")
